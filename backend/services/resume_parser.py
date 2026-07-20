"""Resume text extraction and deterministic career profile parsing."""
from __future__ import annotations

import io
import re
import shutil
import subprocess
import tempfile
from collections import Counter
from pathlib import Path

from docx import Document
from pypdf import PdfReader

SUPPORTED_RESUME_TYPES = {"pdf", "doc", "docx", "txt", "md"}
SKILL_KEYWORDS = [
    "Python", "Java", "JavaScript", "TypeScript", "React", "Vue", "Next.js", "FastAPI",
    "Spring", "MySQL", "PostgreSQL", "Redis", "Kafka", "Docker", "Kubernetes", "AWS",
    "产品规划", "用户研究", "数据分析", "增长", "运营", "销售", "项目管理", "团队管理",
]
METRIC_RE = re.compile(r"(?:提升|增长|降低|减少|节省|达到|完成|负责).{0,40}?(?:\d+(?:\.\d+)?%?|[一二三四五六七八九十百千万]+(?:万|亿)?)")
YEAR_RE = re.compile(r"(?:19|20)\d{2}(?:[./年-]\d{1,2})?")


def extract_resume_text(content: bytes, extension: str) -> str:
    extension = extension.lower().lstrip(".")
    if extension == "pdf":
        reader = PdfReader(io.BytesIO(content))
        extracted = "\n".join((page.extract_text() or "") for page in reader.pages)
        if len(re.sub(r"\s+", "", extracted)) >= 30:
            return extracted
        # Scanned PDFs: render a bounded number of pages and OCR with Tesseract.
        try:
            import fitz
            import pytesseract
            from PIL import Image

            document = fitz.open(stream=content, filetype="pdf")
            ocr_pages = []
            for page_index in range(min(12, document.page_count)):
                page = document.load_page(page_index)
                pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                image = Image.frombytes("RGB", (pixmap.width, pixmap.height), pixmap.samples)
                ocr_pages.append(pytesseract.image_to_string(image, lang="chi_sim+eng"))
            ocr_text = "\n".join(ocr_pages).strip()
            return ocr_text or extracted
        except Exception:
            return extracted
    if extension == "doc":
        antiword = shutil.which("antiword")
        if not antiword:
            raise ValueError("当前环境未安装旧版 Word 解析器")
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as temp:
            temp.write(content)
            temp_path = Path(temp.name)
        try:
            result = subprocess.run([antiword, str(temp_path)], check=True, capture_output=True, timeout=30)
            for encoding in ("utf-8", "gb18030"):
                try:
                    return result.stdout.decode(encoding)
                except UnicodeDecodeError:
                    continue
            return result.stdout.decode("utf-8", errors="replace")
        finally:
            temp_path.unlink(missing_ok=True)
    if extension == "docx":
        document = Document(io.BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                paragraphs.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(paragraphs)
    if extension in {"txt", "md"}:
        for encoding in ("utf-8", "gb18030"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
    raise ValueError("无法解析该简历格式")


def parse_resume_text(text: str) -> dict:
    normalized = re.sub(r"\r\n?", "\n", text)
    lines = [re.sub(r"\s+", " ", line).strip() for line in normalized.splitlines()]
    lines = [line for line in lines if line]
    skills = [keyword for keyword in SKILL_KEYWORDS if keyword.lower() in normalized.lower()]
    metric_lines = [line for line in lines if METRIC_RE.search(line)][:12]
    dated_lines = [line for line in lines if YEAR_RE.search(line)][:20]
    section_names = []
    for line in lines:
        compact = line.replace(" ", "")
        if len(compact) <= 12 and any(key in compact for key in ("教育", "经历", "项目", "技能", "证书", "自我评价", "工作")):
            section_names.append(line)
    return {
        "headline": lines[0][:120] if lines else "",
        "skills": skills,
        "quantified_achievements": metric_lines,
        "timeline_hints": dated_lines,
        "sections": list(dict.fromkeys(section_names))[:12],
        "word_count": len(re.findall(r"[\w\u4e00-\u9fff]+", normalized)),
        "completeness": {
            "has_skills": bool(skills),
            "has_metrics": bool(metric_lines),
            "has_timeline": bool(dated_lines),
            "score": sum((bool(skills), bool(metric_lines), bool(dated_lines))) * 30 + min(10, len(lines) // 5),
        },
    }


def stories_from_resume(parsed: dict, source_resume_id: int | None = None) -> list[dict]:
    stories = []
    for index, achievement in enumerate(parsed.get("quantified_achievements") or []):
        stories.append({
            "title": f"量化成果 {index + 1}",
            "situation": "来自简历的项目或工作经历",
            "task": "明确当时承担的目标和责任",
            "action": achievement,
            "result": achievement,
            "tags": ["量化成果"],
            "metrics": {"source": "resume", "source_resume_id": source_resume_id},
        })
    return stories[:8]
